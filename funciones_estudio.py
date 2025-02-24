import tabula
import pandas as pd
import re
import sqlite3
from pandasql import sqldf
import numpy as np
import warnings
from datetime import datetime
from PyPDF2 import PdfReader
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm  # Importar Mm para definir tamaños en milímetros

warnings.simplefilter(action='ignore', category=Warning)

# Conexión a la base de datos SQLite
Apuntador = sqlite3.connect('DatosApp.db')
cursor = Apuntador.cursor()

def formatear_tupla(array):
  valores = tuple(array)
  if len(valores) == 1:
      valores_str = f"({valores[0]})"  # Sin coma si hay solo un valor
  else:
      valores_str = str(tuple(valores))  # Convierte la lista en tupla (válida en SQL)
  return valores_str

def Obtener_Informacion_PDF(pdf_path):
    # Obtener todas las tablas presentes en el PDF
    tablas = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)

    # Capturar tabla con información básica del estudiante
    for i in range(len(tablas)):
        if len(tablas[i][tablas[i].iloc[:, 0] == "Nombre:"]) > 0:
            info_estudiante = pd.DataFrame(tablas[i].iloc[:, :2].values).dropna()

    # Cargar el archivo PDF como texto
    reader = PdfReader(pdf_path)

    # Extraer texto de todas las páginas
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    # Buscar las secciones requeridas usando expresiones regulares
    creditos_dt = re.search(r"Créditos de estudio doble titulación\s+(\d+)", text).group(1)
    porcentaje_avance = re.search(r"Porcentaje de avance\s+([\d,.]+%)", text).group(1)
    creditos_le_pendientes = re.search(r"Libre Elección\s+\d+\s+\d+\s+(\d+)", text).group(1)
    papa = re.search(r"(\d,\d)\s*\(Acumulado\)\s*PREGRADO\s*P\.A\.P\.A", text).group(1).replace(',', '.')

    nueva_info = np.empty((4,len(info_estudiante.columns)),dtype=object)
    nueva_info[0,0] = "Créditos de estudio doble titulación"
    nueva_info[0,1] = creditos_dt
    nueva_info[1,0] = "Porcentaje de avance"
    nueva_info[1,1] = porcentaje_avance
    nueva_info[2,0] = "Pendientes Libre Elección"
    nueva_info[2,1] = creditos_le_pendientes
    nueva_info[3,0] = "P.A.P.A:"
    nueva_info[3,1] = papa

    info_estudiante = pd.concat([info_estudiante, pd.DataFrame(nueva_info)], ignore_index=True)

    # Construir la tabla con la información de asignaturas
    asignaturas_estudiante = None
    for i in range(len(tablas)):
        if len(tablas[i][tablas[i].iloc[:, 0] == "Asignatura"]) > 0 or "Asignatura" in tablas[i].columns:
            if asignaturas_estudiante is None:
                asignaturas_estudiante = pd.DataFrame(tablas[i].values)
                if "Asignatura" in tablas[i].columns:
                  columnas = tablas[i].columns
                  eliminar_primera = False
                else:
                  eliminar_primera = True
            else:
                asignaturas_estudiante = pd.concat([asignaturas_estudiante, pd.DataFrame(tablas[i].values)])

    asignaturas_estudiante = asignaturas_estudiante.iloc[:,:9]

    # Unir filas fragmentadas
    filas_completas = []
    fila_actual = None

    for _, fila in asignaturas_estudiante.iterrows():
        if pd.notna(fila[0]):  # Si la fila tiene una asignatura, es una nueva fila válida
            if fila_actual is not None:
                filas_completas.append(fila_actual)
            fila_actual = fila.copy()
        elif fila_actual is not None:
            for col in asignaturas_estudiante.columns:
                if pd.isna(fila_actual[col]) and pd.notna(fila[col]):
                    fila_actual[col] = fila[col]

    if fila_actual is not None:
        filas_completas.append(fila_actual)

    asignaturas_estudiante = pd.DataFrame(filas_completas)
    asignaturas_estudiante.reset_index(drop=True, inplace=True)

    # Eliminar encabezados repetidos y asignar columnas
    if eliminar_primera:
      asignaturas_estudiante.columns = asignaturas_estudiante.iloc[0]
      asignaturas_estudiante = asignaturas_estudiante.drop(asignaturas_estudiante.index[0])
    else:
      asignaturas_estudiante.columns = columnas

    # Periodo
    periodo_inicio = info_estudiante.loc[info_estudiante[0] == "Periodo de inicio:"][1].values[0]
    asignaturas_estudiante['Periodo'] = None
    periodo_actual = periodo_inicio

    for i, row in asignaturas_estudiante.iterrows():
        asignatura_str = str(row['Asignatura'])
        if "PERIODO" in asignatura_str:
            match = re.search(r'(\d{4}-\dS)', asignatura_str)
            if match:
                periodo_actual = match.group(1)
        else:
            asignaturas_estudiante.at[i, 'Periodo'] = periodo_actual

    asignaturas_estudiante = asignaturas_estudiante[~asignaturas_estudiante['Asignatura'].str.contains("PERIODO", na=False)]

    # Formatear los datos
    asignaturas_estudiante["Cod_Asignatura"] = asignaturas_estudiante['Asignatura'].str.extract(r'\(([^)]+)\)')
    asignaturas_estudiante["Cod_Asignatura"] = asignaturas_estudiante["Cod_Asignatura"].str.extract(r'(\d+)')
    asignaturas_estudiante['Cod_Asignatura'] = asignaturas_estudiante['Cod_Asignatura'].fillna(method='bfill')
    asignaturas_estudiante['Nom_Asignatura'] = (
        asignaturas_estudiante['Asignatura']
        .str.replace(r'\([^)]*\)', '', regex=True)
        .str.strip()
    )

    # Convertir tipos de datos
    asignaturas_estudiante["Creditos"] = pd.to_numeric(asignaturas_estudiante["Créditos"], errors="coerce").fillna(0).astype(int)
    asignaturas_estudiante["Calificacion"] = (
        asignaturas_estudiante["Calificación"].str.extract(r'(\d+(,\d+)?)')[0]
        .str.replace(",", ".", regex=False)
        .astype(float)
    )

    # Filtrar asignaturas válidas
    asignaturas_estudiante = asignaturas_estudiante[asignaturas_estudiante['Calificacion'].notna()]
    asignaturas_estudiante = asignaturas_estudiante[asignaturas_estudiante['Anulada'] != 'SI']
    asignaturas_estudiante = asignaturas_estudiante[asignaturas_estudiante['Calificacion'] >= 3]
    asignaturas_estudiante = asignaturas_estudiante[asignaturas_estudiante["Tipología"] != "Nivelación (E)"]
    # Trabajo de grado no es equivalible ni convalidable
    asignaturas_estudiante = asignaturas_estudiante[~asignaturas_estudiante["Tipología"].str.contains(r"\(P\)", na=False)]

    # Seleccionar columnas finales
    asignaturas_estudiante = asignaturas_estudiante[['Periodo', 'Cod_Asignatura', 'Nom_Asignatura', 'Creditos', 'Calificacion','Tipología']]

    return info_estudiante,asignaturas_estudiante


# Queremos obtener una lista de EQUIVALENCIAS CANDIDATAS

def Generar_Lista_Candidatas(asignaturas_estudiante):
    # Query para obtener las asignaturas que cursó y que pertenecen al programa
    # Esto es Equivalencias con igual código
    lista_asignaturas_cursadas = tuple(asignaturas_estudiante["Cod_Asignatura"])
    query_1 = f"""
        SELECT
            Cod_Asignatura_CC AS Código,
            Nom_Agrupacion AS Agrupación,
            Tipo
        FROM Asignaturas_CC
        WHERE Cod_Asignatura_CC IN {formatear_tupla(lista_asignaturas_cursadas)}"""
    asignaturas_iguales = pd.read_sql_query(query_1, Apuntador)

    # Join para completar la información de las asignaturas
    query_2 = """
        SELECT
            asignaturas_estudiante.Periodo AS "Periodo Académico",
            asignaturas_estudiante.Cod_Asignatura AS "Código",
            asignaturas_estudiante.Nom_Asignatura AS "Asignatura",
            asignaturas_iguales.Código AS "Código_CC",
            asignaturas_estudiante.Nom_Asignatura AS "Asignatura_CC",
            asignaturas_iguales.Agrupación AS "Agrupación",
            asignaturas_estudiante.Calificacion AS "Nota",
            asignaturas_iguales.Tipo AS "Tipo",
            asignaturas_estudiante.Creditos AS "Créditos"
        FROM asignaturas_estudiante
        JOIN asignaturas_iguales
        ON asignaturas_iguales.Código = asignaturas_estudiante.Cod_Asignatura;
    """

    equivalencias_candidatas = pd.DataFrame(sqldf(query_2))


    # Equivalencias con diferente código
    # Obtener la equivalencia desde la base de datos
    query_3 = f"""
        SELECT
            Cod_Asignatura AS Código,
            Cod_Asignatura_CC AS 'Código_CC'
        FROM Equivalencias
        WHERE Cod_Asignatura IN {formatear_tupla(lista_asignaturas_cursadas)}"""
    asignaturas_equivalentes = pd.read_sql_query(query_3, Apuntador)

    # Join para obtener la agrupación y tipo de las asignaturas equivalentes
    query_4 = f"""
        SELECT
            Cod_Asignatura_CC AS 'Código_CC',
            Nom_Agrupacion AS Agrupación,
            Tipo
        FROM Asignaturas_CC
        WHERE Cod_Asignatura_CC IN {formatear_tupla(asignaturas_equivalentes["Código_CC"])}"""
    info_equivalentes = pd.read_sql_query(query_4, Apuntador)

    # Join para obtener el número de créditos y nombre de las asignaturas equivalentes
    query_5 = f"""
        SELECT
            Cod_Asignatura AS Código_CC,
            Nom_Asignatura AS Asignatura_CC,
            Creditos
        FROM Asignaturas_Info
        WHERE Cod_Asignatura IN {formatear_tupla(asignaturas_equivalentes["Código_CC"])}"""
    info_equivalentes_2 = pd.read_sql_query(query_5, Apuntador)

    # Join para completar la información de las asignaturas
    query_6 = """
    SELECT
        e.Periodo AS "Periodo Académico",
        e.Cod_Asignatura AS "Código",
        e.Nom_Asignatura AS "Asignatura",
        a.Código_CC AS "Código_CC",
        ie2.Asignatura_CC AS "Asignatura_CC",
        ie.Agrupación AS "Agrupación",
        e.Calificacion AS "Nota",
        ie.Tipo AS "Tipo",
        ie2.Creditos AS "Créditos"
    FROM asignaturas_estudiante AS e
    JOIN asignaturas_equivalentes AS a
    ON e.Cod_Asignatura = a.Código
    JOIN info_equivalentes AS ie
    ON a.Código_CC = ie.Código_CC
    JOIN info_equivalentes_2 AS ie2
    ON ie.Código_CC = ie2.Código_CC;
    """
    # Concatenar las equivalencias de asignaturas iguales con las de diferentes
    equivalencias_candidatas = pd.concat([equivalencias_candidatas,pd.DataFrame(sqldf(query_6))],axis=0)
    return equivalencias_candidatas

def Generar_Estudio(info_estudiante,asignaturas_estudiante,lista_candidatas):
    # Consulta de agrupaciones
    query_7 = 'SELECT * FROM Agrupaciones_CC'
    agrupaciones = pd.read_sql_query(query_7, Apuntador)
    agrupaciones

    # Cursadas en Fundamentación (Por el momento únicamente las obligatorias)
    cursadas_fundamentacion = lista_candidatas[lista_candidatas["Tipo"] == "B"]

    # Pendientes en Fundamentación - Obligatoria B
    query_8 = f"""
    SELECT
        Cod_Asignatura_CC AS 'Código',
        Nom_Agrupacion AS 'Agrupación'
    FROM Asignaturas_CC
    WHERE Cod_Asignatura_CC NOT IN {formatear_tupla(cursadas_fundamentacion["Código_CC"])} AND Tipo = 'B'
    """
    pendientes_b = pd.read_sql_query(query_8, Apuntador)

    query_9 = f"""
    SELECT
        Cod_Asignatura as 'Código',
        Nom_Asignatura as 'Asignatura',
        Creditos as 'Créditos'
    FROM Asignaturas_Info
    WHERE Cod_Asignatura IN {formatear_tupla(pendientes_b["Código"])}
    """
    pendientes_b_info = pd.read_sql_query(query_9, Apuntador)
    pendientes_b_info

    query_10 = f"""
    SELECT
        p.Agrupación AS 'Agrupación',
        i.Código as 'Código',
        i.Asignatura as 'Asignatura',
        i.Créditos as 'Créditos'
    FROM pendientes_b_info as i
    JOIN pendientes_b as p
    ON i.Código = p.Código
    """
    tabla_pendientes_b = pd.DataFrame(sqldf(query_10)).sort_values(by="Agrupación")

    # Calcular la suma de créditos por agrupación
    suma_creditos = tabla_pendientes_b.groupby("Agrupación")["Créditos"].sum().reset_index()
    suma_creditos.rename(columns={"Créditos": "Créditos pendientes por cursar de la agrupación"}, inplace=True)

    # Unir esta información al DataFrame original
    tabla_pendientes_b = tabla_pendientes_b.merge(suma_creditos, on="Agrupación", how="left")

    # Pendientes en Disciplinar - Obligatoria C y Trabajo de Grado P
    # Cursadas en Fundamentación (Por el momento únicamente las obligatorias)
    cursadas_disciplinar = lista_candidatas[lista_candidatas["Tipo"] == "C"]

    query_11 = f"""
    SELECT
        Cod_Asignatura_CC AS 'Código',
        Nom_Agrupacion AS 'Agrupación'
    FROM Asignaturas_CC
    WHERE Cod_Asignatura_CC NOT IN {formatear_tupla(cursadas_disciplinar["Código_CC"])} AND Tipo = 'C'
    """
    pendientes_c = pd.read_sql_query(query_11, Apuntador)

    query_12 = f"""
    SELECT
        Cod_Asignatura as 'Código',
        Nom_Asignatura as 'Asignatura',
        Creditos as 'Créditos'
    FROM Asignaturas_Info
    WHERE Cod_Asignatura IN {formatear_tupla(pendientes_c["Código"])}
    """
    pendientes_c_info = pd.read_sql_query(query_12, Apuntador)
    pendientes_c_info

    query_13 = f"""
    SELECT
        p.Agrupación AS 'Agrupación',
        i.Código as 'Código',
        i.Asignatura as 'Asignatura',
        i.Créditos as 'Créditos'
    FROM pendientes_c_info as i
    JOIN pendientes_c as p
    ON i.Código = p.Código
    """
    tabla_pendientes_c = pd.DataFrame(sqldf(query_13)).sort_values(by="Agrupación")

    # Calcular la suma de créditos por agrupación
    suma_creditos = tabla_pendientes_c.groupby("Agrupación")["Créditos"].sum().reset_index()
    suma_creditos.rename(columns={"Créditos": "Créditos pendientes por cursar de la agrupación"}, inplace=True)

    # Unir esta información al DataFrame original
    tabla_pendientes_c = tabla_pendientes_c.merge(suma_creditos, on="Agrupación", how="left")


    # Pendientes en Disciplinar - Obligatoria P Trabajo de Grado
    query_14 = f"""
    SELECT
        a.Nom_Agrupacion AS 'Agrupación',
        a.Cod_Asignatura_CC AS 'Código',
        i.Nom_Asignatura AS 'Asignatura',
        i.Creditos AS 'Créditos',
        agrup.Cant_Optativos AS 'Créditos pendientes por cursar de la agrupación'
    FROM Asignaturas_CC AS a
    JOIN Asignaturas_Info AS i
    ON i.Cod_Asignatura = a.Cod_Asignatura_CC
    JOIN Agrupaciones_CC AS agrup
    ON agrup.Nom_Agrupacion = a.Nom_Agrupacion
    WHERE a.Nom_Agrupacion = 'TRABAJO DE GRADO'
    """
    tabla_pendientes_disciplinar = pd.concat([tabla_pendientes_c,pd.read_sql_query(query_14, Apuntador)]).reset_index(drop=True)

    # Contador de créditos Optativas O y T
    creditos_contados = pd.DataFrame(columns=["Agrupación","Tipo","Créditos aprobados","Créditos excedentes"])

    # Seleccionar en Fundamentación Optativa - O las asignaturas a equivaler
    # Agrupación de fundamentación
    lista_agrup_fund = agrupaciones[agrupaciones["Tipologia"]=="Fundamentación"]["Nom_Agrupacion"]

    # Seleccionar las asignaturas con mayor nota hasta completar la agrupación
    for agrupacion in lista_agrup_fund:
        cantidad_creditos_o = agrupaciones[agrupaciones["Nom_Agrupacion"]==agrupacion]["Cant_Optativos"].values[0]
        asignaturas = lista_candidatas[(lista_candidatas["Agrupación"] == agrupacion) & (lista_candidatas["Tipo"] == "O")].sort_values(by="Nota", ascending=False).reset_index(drop=True)
        suma_creditos = 0
        i = 0
        while suma_creditos < cantidad_creditos_o and 0<len(asignaturas):
            # Seleccionar el primer registro - el que tiene mejor nota
            asignatura = asignaturas.iloc[0]

            # Agregarlo a cursadas_fundamentacion
            cursadas_fundamentacion = pd.concat([cursadas_fundamentacion, pd.DataFrame([asignatura])], ignore_index=True)
            suma_creditos = suma_creditos + asignatura["Créditos"]

            # Eliminar el primer registro de tabla
            asignaturas = asignaturas.iloc[1:].reset_index(drop=True)

            # Incrementar el contador
            i += 1
        creditos_contados = pd.concat([creditos_contados, pd.DataFrame([[agrupacion,"Fundamentación", suma_creditos, max(suma_creditos-cantidad_creditos_o,0)]], columns=["Agrupación", "Tipo", "Créditos aprobados","Créditos excedentes"])], ignore_index=True)


    # Seleccionar en Disciplinar - T las asignaturas a equivaler
    # Agrupación de Disciplinar
    lista_agrup_disciplinar = agrupaciones[agrupaciones["Tipologia"]=="Disciplinar"]["Nom_Agrupacion"]

    # Seleccionar las asignaturas con mayor nota hasta completar la agrupación
    for agrupacion in lista_agrup_disciplinar:
        cantidad_creditos_t = agrupaciones[agrupaciones["Nom_Agrupacion"]==agrupacion]["Cant_Optativos"].values[0]
        asignaturas = lista_candidatas[(lista_candidatas["Agrupación"] == agrupacion) & (lista_candidatas["Tipo"] == "T")].sort_values(by="Nota", ascending=False).reset_index(drop=True)
        suma_creditos = 0
        i = 0
        while suma_creditos < cantidad_creditos_t and 0<len(asignaturas):
            # Seleccionar el primer registro - el que tiene mejor nota
            asignatura = asignaturas.iloc[0]

            # Agregarlo a cursadas_fundamentacion
            cursadas_disciplinar = pd.concat([cursadas_disciplinar, pd.DataFrame([asignatura])], ignore_index=True)
            suma_creditos = suma_creditos + int(asignatura["Créditos"])

            # Eliminar el primer registro de tabla
            asignaturas = asignaturas.iloc[1:].reset_index(drop=True)

            # Incrementar el contador
            i += 1
        creditos_contados = pd.concat([creditos_contados, pd.DataFrame([[agrupacion, "Disciplinar", suma_creditos, max(suma_creditos-cantidad_creditos_t,0)]], columns=["Agrupación","Tipo","Créditos aprobados","Créditos excedentes"])], ignore_index=True)

    # Visualizar pendientes en Fundamentación Optativa O
    # Pendientes en Fundamentación - Optativa O
    pendientes_o = agrupaciones[agrupaciones["Tipologia"]=="Fundamentación"][["Nom_Agrupacion", "Cant_Optativos"]]
    pendientes_o.rename(columns={"Nom_Agrupacion": "Agrupación", "Cant_Optativos": "Créditos exigidos en la agrupación"}, inplace=True)
    pendientes_o["Créditos pendientes por cursar de la agrupación"] = (
        pendientes_o["Créditos exigidos en la agrupación"] -
        creditos_contados[creditos_contados["Agrupación"].isin(pendientes_o["Agrupación"])]["Créditos aprobados"].values +
        creditos_contados[creditos_contados["Agrupación"].isin(pendientes_o["Agrupación"])]["Créditos excedentes"].values
    ).clip(lower=0)

    # Visualizar pendientes en Disciplinar Optativa T
    # Pendientes en Disciplinar - Optativa T
    pendientes_t = agrupaciones[agrupaciones["Tipologia"]=="Disciplinar"][["Nom_Agrupacion", "Cant_Optativos"]]
    pendientes_t.rename(columns={"Nom_Agrupacion": "Agrupación", "Cant_Optativos": "Créditos exigidos en la agrupación"}, inplace=True)
    pendientes_t["Créditos pendientes por cursar de la agrupación"] = (
        pendientes_t["Créditos exigidos en la agrupación"] -
        creditos_contados[creditos_contados["Agrupación"].isin(pendientes_t["Agrupación"])]["Créditos aprobados"].values +
        creditos_contados[creditos_contados["Agrupación"].isin(pendientes_t["Agrupación"])]["Créditos excedentes"].values
    ).clip(lower=0)

    # Seleccionar en Libre Elección - L las asignaturas a equivaler
    candidatas_libre = asignaturas_estudiante[~(asignaturas_estudiante["Cod_Asignatura"].isin(lista_candidatas["Código"]))]
    # Ordenar por calificación
    candidatas_libre = candidatas_libre.sort_values(by="Calificacion", ascending=False).reset_index(drop=True)

    # Seleccionar en Libre Elección - L las asignaturas a equivaler
    creditos_libre = agrupaciones[agrupaciones["Tipologia"]=="Libre Elección"]["Cant_Optativos"].values[0]
    cursadas_libre = pd.DataFrame(columns=candidatas_libre.columns)
    suma_creditos = 0
    i = 0
    while suma_creditos<creditos_libre and len(candidatas_libre)>0:
        # Seleccionar el primer registro - el que tiene mejor nota
        asignatura = candidatas_libre.iloc[0]

        # Agregarlo a cursadas_libre
        cursadas_libre = pd.concat([cursadas_libre, pd.DataFrame([asignatura])], ignore_index=True)
        suma_creditos = suma_creditos + int(asignatura["Creditos"])

        # Eliminar el primer registro de tabla
        candidatas_libre = candidatas_libre.iloc[1:].reset_index(drop=True)

        # Incrementar el contador
        i += 1

    #Guardar dato de los excedentes
    creditos_contados = pd.concat([creditos_contados, pd.DataFrame([["LIBRE ELECCIÓN", "Libre Elección", suma_creditos, max(suma_creditos-creditos_libre,0)]], columns=["Agrupación","Tipo","Créditos aprobados","Créditos excedentes"])], ignore_index=True)

    libre_cursadas = pd.DataFrame(columns=["Periodo Académico", "Código", "Asignatura", "Código_CC", "Asignatura_CC", "Nota", "#Créditos"])
    libre_cursadas.iloc[:,0] = cursadas_libre["Periodo"]
    libre_cursadas.iloc[:,1] = cursadas_libre["Cod_Asignatura"]
    libre_cursadas.iloc[:,2] = cursadas_libre["Nom_Asignatura"]
    libre_cursadas.iloc[:,3] = cursadas_libre["Cod_Asignatura"]
    libre_cursadas.iloc[:,4] = cursadas_libre["Nom_Asignatura"]
    libre_cursadas.iloc[:,5] = cursadas_libre["Calificacion"]
    libre_cursadas.iloc[:,6] = cursadas_libre["Creditos"]

    #Tabla de Resumen General
    resumen_general = pd.DataFrame(columns=["Créditos","Fund. Obl. B","Fund. Opt. O","Disc. Obl. C","Disc. Opt. T","Trabajo de Grado P","Libre Elección L", "Total"])
    resumen_general["Créditos"] = ["Exigidos","Convalidados/equivalentes","Pendientes","Excedentes"]

    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Fund. Obl. B"] = agrupaciones[agrupaciones["Tipologia"]=="Fundamentación"]["Cant_Obligatorios"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Fund. Opt. O"] = agrupaciones[agrupaciones["Tipologia"]=="Fundamentación"]["Cant_Optativos"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Disc. Obl. C"] = agrupaciones[agrupaciones["Tipologia"]=="Disciplinar"]["Cant_Obligatorios"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Disc. Opt. T"] = agrupaciones[agrupaciones["Tipologia"]=="Disciplinar"]["Cant_Optativos"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Trabajo de Grado P"] = agrupaciones[agrupaciones["Tipologia"]=="Trabajo de grado"]["Cant_Optativos"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Libre Elección L"] = agrupaciones[agrupaciones["Tipologia"]=="Libre Elección"]["Cant_Optativos"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Total"] = resumen_general.loc[resumen_general["Créditos"] == "Exigidos", resumen_general.columns[1:]].sum(axis=1)

    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Fund. Obl. B"] = cursadas_fundamentacion[cursadas_fundamentacion["Tipo"]=="B"]["Créditos"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Fund. Opt. O"] = cursadas_fundamentacion[cursadas_fundamentacion["Tipo"]=="O"]["Créditos"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Disc. Obl. C"] = cursadas_disciplinar[cursadas_disciplinar["Tipo"]=="C"]["Créditos"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Disc. Opt. T"] = cursadas_disciplinar[cursadas_disciplinar["Tipo"]=="T"]["Créditos"].astype(int).sum()
    # De acuerdo a la norma el trabajo de grado no es equivalible ni convalidable
    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Trabajo de Grado P"] = 0
    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Libre Elección L"] = libre_cursadas["#Créditos"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Total"] = resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", resumen_general.columns[1:]].sum(axis=1)


    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Fund. Obl. B"] = tabla_pendientes_b["Créditos"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Disc. Obl. C"] = tabla_pendientes_c["Créditos"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Fund. Opt. O"] = pendientes_o["Créditos pendientes por cursar de la agrupación"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Disc. Opt. T"] = pendientes_t["Créditos pendientes por cursar de la agrupación"].astype(int).sum()
    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Trabajo de Grado P"] = agrupaciones[agrupaciones["Tipologia"]=="Trabajo de grado"]["Cant_Optativos"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Libre Elección L"] = max(0,resumen_general.loc[resumen_general["Créditos"] == "Exigidos", "Libre Elección L"].values[0] - resumen_general.loc[resumen_general["Créditos"] == "Convalidados/equivalentes", "Libre Elección L"].values[0])
    resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Total"] = resumen_general.loc[resumen_general["Créditos"] == "Pendientes", resumen_general.columns[1:]].sum(axis=1)

    # En las obligatorias nunca se obtienen créditos excedentes
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Fund. Obl. B"] = 0
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Disc. Obl. C"] = 0
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Trabajo de Grado P"] = 0
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Fund. Opt. O"] = creditos_contados[creditos_contados["Tipo"]=="Fundamentación"]["Créditos excedentes"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Disc. Opt. T"] = creditos_contados[creditos_contados["Tipo"]=="Disciplinar"]["Créditos excedentes"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Libre Elección L"] = creditos_contados[creditos_contados["Tipo"]=="Libre Elección"]["Créditos excedentes"].sum()
    resumen_general.loc[resumen_general["Créditos"] == "Excedentes", "Total"] = resumen_general.loc[resumen_general["Créditos"] == "Excedentes", resumen_general.columns[1:]].sum(axis=1)

    return resumen_general, cursadas_fundamentacion, cursadas_disciplinar, libre_cursadas, tabla_pendientes_b, pendientes_o, tabla_pendientes_disciplinar, pendientes_t

def Exportar_Estudio(info_estudiante, resumen_general, cursadas_fundamentacion, cursadas_disciplinar, libre_cursadas, tabla_pendientes_b, pendientes_o, tabla_pendientes_disciplinar, pendientes_t):
    # Crear el documento Word
    doc = Document()

    # Configurar tamaño de papel A3 en orientación vertical (portrait)
    section = doc.sections[0]
    section.page_width = Mm(297)   # Ancho para A3 en portrait
    section.page_height = Mm(420)  # Alto para A3 en portrait

    # Función para establecer bordes negros en una tabla
    def set_table_border(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for border in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tblBorders.append(element)
        tblPr.append(tblBorders)

    # Función que inserta un DataFrame como tabla en el documento
    def add_table_from_df(document, df, title=None):
        if title:
            document.add_heading(title, level=2)
        # Se crea la tabla con filas = cabecera + datos y columnas según el DataFrame
        table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')

        # Agregar la fila de cabecera
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = str(col)

        # Agregar las filas de datos
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i+1, j).text = str(df.iloc[i, j])

        # Aseguramos que los bordes sean negros
        set_table_border(table)
        document.add_paragraph("")  # Línea en blanco para separar
        return table

    # Información inicial
    query_15 = """
    SELECT *
    FROM Planes_de_Estudio"""
    info_plan = pd.read_sql_query(query_15, Apuntador)

    doc.add_heading("1. Datos Generales", level=1)
    info_inicial = pd.DataFrame([
        ["Nombre", info_estudiante.loc[info_estudiante[0] == "Nombre:"][1].values[0]],
        ["DNI", info_estudiante.loc[info_estudiante[0] == "Documento:"][1].values[0]],
        ["Plan de estudios origen (1er plan):", info_estudiante.loc[info_estudiante[0] == "Plan:"][1].values[0]],
        ["Plan de estudios doble titulación (2° plan) - Sede",
        str(info_plan["Codigo_plan"].values[0]) + " - " + str(info_plan["Nombre_plan"].values[0])],
        ["Fecha de la solicitud", datetime.today().strftime("%d/%m/%Y")]
    ])
    add_table_from_df(doc, info_inicial)

    doc.add_heading("2. Información académica", level=1)
    porcentaje_avance = info_estudiante.loc[info_estudiante[0] == "Porcentaje de avance"][1].values[0]
    porcentaje_avance_num = float(porcentaje_avance.split("%")[0].replace(',', '.'))
    porcentaje_avance_40 = "SI" if porcentaje_avance_num >= 40 else "NO"
    papa = info_estudiante.loc[info_estudiante[0] == "P.A.P.A:"][1].values[0]
    creditos_dt = info_estudiante.loc[info_estudiante[0] == "Créditos de estudio doble titulación"][1].values[0]

    info_academica = pd.DataFrame([
        ["¿Tuvo calidad de estudiante en el 2° plan?", "NO"],
        ["¿Se encuentra matriculado o en reserva de cupo al momento de presentar la solicitud?", "SI"],
        ["Porcentaje de avance en el plan de estudios origen", porcentaje_avance],
        ["¿El estudiante ha cursado el 40% de los créditos de su primer plan de estudios?", porcentaje_avance_40],
        ["P.A.P.A. en el primer plan de estudio", papa],
        ["Créditos estudio doble titulación (información reportada en el resumen de créditos del SIA)", creditos_dt]
    ])
    add_table_from_df(doc, info_academica)

    doc.add_heading("3. Resumen general de créditos del segundo plan de estudios:", level=1)
    add_table_from_df(doc, resumen_general)

    doc.add_heading("4. Cuadro de equivalencia y convalidaciones de asignaturas", level=1)

    doc.add_heading("Componente de FUNDAMENTACIÓN", level=2)
    cursadas_fundamentacion["Créditos_B"] = cursadas_fundamentacion.apply(lambda row: row["Créditos"] if row["Tipo"] == "B" else "", axis=1)
    cursadas_fundamentacion["Créditos_O"] = cursadas_fundamentacion.apply(lambda row: row["Créditos"] if row["Tipo"] == "O" else "", axis=1)
    cursadas_fundamentacion.drop(columns=["Tipo", "Créditos"], inplace=True)
    add_table_from_df(doc, cursadas_fundamentacion)

    doc.add_heading("Componente DISCIPLINAR", level=2)
    cursadas_disciplinar["Créditos_C"] = cursadas_disciplinar.apply(lambda row: row["Créditos"] if row["Tipo"] == "C" else "", axis=1)
    cursadas_disciplinar["Créditos_T"] = cursadas_disciplinar.apply(lambda row: row["Créditos"] if row["Tipo"] == "T" else "", axis=1)
    cursadas_disciplinar.drop(columns=["Tipo", "Créditos"], inplace=True)
    add_table_from_df(doc, cursadas_disciplinar)

    doc.add_heading("Componente de LIBRE ELECCIÓN", level=2)
    add_table_from_df(doc, libre_cursadas)

    doc.add_heading("5. Asignaturas pendientes por cursar en el segundo plan de estudios", level=1)
    doc.add_paragraph("Pendientes en Fundamentación - Obligatoria B")
    add_table_from_df(doc, tabla_pendientes_b)
    doc.add_paragraph("Pendientes en Fundamentación - Optativa O")
    add_table_from_df(doc, pendientes_o)
    doc.add_paragraph("Pendientes en Disciplinar - Obligatoria C y P")
    add_table_from_df(doc, tabla_pendientes_disciplinar)
    doc.add_paragraph("Pendientes en Disciplinar - Optativa T")
    add_table_from_df(doc, pendientes_t)

    add_table_from_df(doc,pd.DataFrame(columns=["Componente de Libre Elección (L) (Créditos pendientes)",resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Libre Elección L"].values[0]]))

    doc.add_heading("6. El Consejo de Facultad Recomienda / No recomienda", level=1)
    recomienda = None
    razon = ""
    if porcentaje_avance_num < 40:
        recomienda = False
        razon = "No dispone de un porcentaje de avance mayor al 40 en el primer plan de estudios."
    elif float(papa) >= 4.3:
        recomienda = True
        if int(creditos_dt) >= int(resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Total"].values[0]):
          razon = "Dispone de los créditos suficientes y además presenta un PAPA mayor o igual a 4.3"
        else:
          razon = ("Aunque no dispone de los créditos suficientes para cursar las asignaturas pendientes de "
                "aprobación en el segundo plan, presenta un PAPA mayor o igual a 4,3.")
    elif int(resumen_general.loc[resumen_general["Créditos"] == "Pendientes", "Total"].values[0]) > int(creditos_dt):
        recomienda = False
        razon = "No dispone de los créditos suficientes para cursar las asignaturas pendientes de aprobación en el segundo plan."
    else:
        recomienda = True
        razon = ("Dispone de los créditos suficientes para cursar las asignaturas pendientes de aprobación "
                "en el segundo plan de estudios.")

    doc.add_paragraph("Recomendación: " + ("Recomienda" if recomienda else "No recomienda"))
    doc.add_paragraph(razon)

    # Guardar el documento Word
    nombre_estudiante = info_estudiante.loc[info_estudiante[0] == "Nombre:"][1].values[0]
    doc.save(f"Estudios/{nombre_estudiante}_estudio_dt.docx")

def Realizar_Estudio(pdf_path):
  info_estudiante, asignaturas_estudiante = Obtener_Informacion_PDF(pdf_path)
  lista_candidatas = Generar_Lista_Candidatas(asignaturas_estudiante)
  resumen_general, cursadas_fundamentacion, cursadas_disciplinar, libre_cursadas, tabla_pendientes_b, pendientes_o, tabla_pendientes_disciplinar, pendientes_t = Generar_Estudio(info_estudiante,asignaturas_estudiante,lista_candidatas)
  Exportar_Estudio(info_estudiante,resumen_general, cursadas_fundamentacion, cursadas_disciplinar, libre_cursadas, tabla_pendientes_b, pendientes_o, tabla_pendientes_disciplinar, pendientes_t)

def Actualizar_Historia(pdf_plan_origen,pdf_plan_cc):
  info_estudiante_origen, asignaturas_estudiante_origen = Obtener_Informacion_PDF(pdf_plan_origen)
  info_estudiante_cc, asignaturas_estudiante_cc = Obtener_Informacion_PDF(pdf_plan_cc)

  lista_candidatas_origen = Generar_Lista_Candidatas(asignaturas_estudiante_origen)

  resumen_general_origen, cursadas_fund_origen, cursadas_disc_origen, libre_cursadas_origen, pendientes_b_origen, pendientes_o_origen, pendientes_disc_origen, pendientes_t_origen = Generar_Estudio(info_estudiante_origen, asignaturas_estudiante_origen, lista_candidatas_origen)

  # Obligatorias B o C
  asignaturas_nuevas = lista_candidatas_origen[(lista_candidatas_origen["Tipo"]=="B") | (lista_candidatas_origen["Tipo"]=="C")]
  query_16 = f"""
          SELECT *
          FROM asignaturas_nuevas
          WHERE Código_CC NOT IN {formatear_tupla(asignaturas_estudiante_cc["Cod_Asignatura"])}"""
  asignaturas_nuevas = pd.DataFrame(sqldf(query_16))

  # Libre Elección L
  pendientes_le = int(info_estudiante_cc.loc[info_estudiante_cc[0] == "Pendientes Libre Elección"][1].values[0])
  if pendientes_le>0:
    libre_asignaturas = libre_cursadas_origen
    libre_asignaturas["Tipo"] = "L"
    libre_asignaturas["Agrupación"] = "LIBRE ELECCIÓN"
    query_17 = f"""
          SELECT *
          FROM libre_asignaturas
          WHERE Código_CC NOT IN {formatear_tupla(asignaturas_estudiante_cc["Cod_Asignatura"])}"""
    asignaturas = pd.DataFrame(sqldf(query_17)).sort_values(by="Créditos").reset_index(drop=True)
    suma_creditos = 0
    i = 0
    while suma_creditos<pendientes_le and len(asignaturas)>0:
            asignatura = asignaturas.iloc[0]

            # Agregarlo a asignaturas_nuevas
            asignaturas_nuevas = pd.concat([asignaturas_nuevas, pd.DataFrame([asignatura])], ignore_index=True)
            suma_creditos = suma_creditos + int(pd.DataFrame([asignatura])["Créditos"])

            # Eliminar el primer registro de tabla
            asignaturas = asignaturas[1:].reset_index(drop=True)

            # Incrementar el contador
            i += 1

  # Optativas O y T
  posibles_optativas_nuevas = lista_candidatas_origen[(lista_candidatas_origen["Tipo"]=="O") | (lista_candidatas_origen["Tipo"]=="T")]
  query_18 = f"""
          SELECT *
          FROM posibles_optativas_nuevas
          WHERE Código_CC NOT IN {formatear_tupla(asignaturas_estudiante_cc["Cod_Asignatura"])}"""
  posibles_optativas_nuevas = pd.DataFrame(sqldf(query_18))

  asignaturas_estudiante_cc["Tipo"] = asignaturas_estudiante_cc["Tipología"].apply(lambda x: re.search(r"\((.*?)\)", x).group(1) if re.search(r"\((.*?)\)", x) else np.nan)
  optativas_cc = asignaturas_estudiante_cc[(asignaturas_estudiante_cc["Tipo"]=="O") | (asignaturas_estudiante_cc["Tipo"]=="T")]

  query_19 = f"""
          SELECT
              Cod_Asignatura_CC as Cod_Asignatura,
              Nom_Agrupacion
          FROM Asignaturas_CC
          WHERE Cod_Asignatura_CC IN {formatear_tupla(optativas_cc["Cod_Asignatura"])}"""
  agrupaciones_optativas = pd.read_sql_query(query_19, Apuntador)

  query_20 = f"""
          SELECT
              optativas_cc.Periodo as 'Periodo Académico',
              optativas_cc.Cod_Asignatura as 'Código',
              optativas_cc.Nom_Asignatura as 'Asignatura',
              agrupaciones_optativas.Cod_Asignatura as 'Código_CC',
              optativas_cc.Nom_Asignatura as 'Asignatura_CC',
              agrupaciones_optativas.Nom_Agrupacion as 'Agrupación',
              optativas_cc.Calificacion as 'Nota',
              optativas_cc.Tipo as 'Tipo',
              optativas_cc.Creditos as 'Créditos'
          FROM optativas_cc
          JOIN agrupaciones_optativas
          ON optativas_cc.Cod_Asignatura=agrupaciones_optativas.Cod_Asignatura;"""
  optativas_cc = pd.DataFrame(sqldf(query_20))

  # Consulta de agrupaciones
  query_20 = 'SELECT * FROM Agrupaciones_CC'
  agrupaciones = pd.read_sql_query(query_20, Apuntador)
  agrupaciones = agrupaciones[(agrupaciones["Nom_Agrupacion"] != "TRABAJO DE GRADO") & (agrupaciones["Nom_Agrupacion"] != "LIBRE ELECCIÓN")]
  for agrupacion in agrupaciones["Nom_Agrupacion"]:
      optativas_agrup = optativas_cc[optativas_cc["Agrupación"] == agrupacion]
      if optativas_agrup["Créditos"].sum() < int(agrupaciones.loc[agrupaciones["Nom_Agrupacion"]==agrupacion]["Cant_Optativos"]):
        asignaturas = posibles_optativas_nuevas[posibles_optativas_nuevas["Agrupación"]==agrupacion].reset_index(drop=True)
        suma_creditos = 0
        i = 0
        pendientes_agrupacion = int(agrupaciones.loc[agrupaciones["Nom_Agrupacion"]==agrupacion]["Cant_Optativos"]) - optativas_agrup["Créditos"].sum()
        while suma_creditos<pendientes_agrupacion and len(asignaturas)>0:
            # Seleccionar el primer registro - el que tiene mejor nota
            asignatura = asignaturas.iloc[0]

            # Agregarlo a asignaturas_nuevas
            asignaturas_nuevas = pd.concat([asignaturas_nuevas, pd.DataFrame([asignatura])], ignore_index=True)
            suma_creditos = suma_creditos + int(pd.DataFrame([asignatura])["Créditos"])

            # Eliminar el primer registro de tabla
            asignaturas = asignaturas[1:].reset_index(drop=True)

            # Incrementar el contador
            i += 1
  return asignaturas_nuevas